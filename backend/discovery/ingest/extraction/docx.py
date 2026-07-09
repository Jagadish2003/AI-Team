"""
R18-A1 / T2 — Word (.docx) text extraction handler.

Extracts paragraph and table text from a Word document with ``python-docx``. An
encrypted document (delivered as an OLE compound file, not a ZIP) is a loud
:data:`~discovery.ingest.extraction.ENCRYPTED` skip (AC3); a corrupt/unparseable
one raises :class:`~discovery.ingest.extraction.ExtractionError` so the ingestor
isolates it per-file (AC5). Table cell text is included so structured content in a
document is not lost. The legacy binary ``.doc`` format is not OOXML and is out of
phase-one scope — it is not routed here.
"""
from __future__ import annotations

import io
import logging

from . import (
    ENCRYPTED,
    NO_HANDLER,
    ExtractedText,
    ExtractionError,
    ExtractionSkipped,
    ExtractionOutcome,
    register_handler,
)
from ._office_common import looks_encrypted_ooxml

logger = logging.getLogger(__name__)


def extract_docx(raw: bytes, fmt: str) -> ExtractionOutcome:
    """Extract paragraph + table text from a .docx, or record a loud skip."""
    if looks_encrypted_ooxml(raw):
        return ExtractionSkipped(ENCRYPTED, "Word document is password-protected")

    try:
        import docx  # python-docx
    except ImportError:  # pragma: no cover - python-docx ships in requirements
        return ExtractionSkipped(NO_HANDLER, "python-docx is not installed")

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:  # noqa: BLE001 — a parse failure is a corrupt file (AC5)
        raise ExtractionError(f"could not parse .docx: {type(exc).__name__}") from exc

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    table_count = 0
    for table in document.tables:
        table_count += 1
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    content = "\n".join(parts).strip()
    return ExtractedText(
        content=content,
        chunk_content_type="prose",
        structure_hints={
            "format": "docx",
            "paragraphs": len(document.paragraphs),
            "tables": table_count,
            "chars": len(content),
        },
        provenance=_docx_provenance(document),
    )


def _docx_provenance(document) -> dict:
    """Best-effort title/author from the document core properties (never fatal)."""
    prov: dict = {}
    try:
        props = document.core_properties
        if props.title:
            prov["title"] = str(props.title)
        if props.author:
            prov["author"] = str(props.author)
    except Exception:  # noqa: BLE001 — core properties are advisory only
        pass
    return prov


register_handler("docx", extract_docx)
